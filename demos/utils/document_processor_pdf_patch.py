import os
from pathlib import Path
from pypdf import PdfReader
from fairlib.core.types import Document

class DocumentProcessor:
    """
    Patched DocumentProcessor with full PDF support via pypdf.
    Extracts text from: .pdf, .txt, .md, .docx (basic), etc.
    Splits into semantic chunks for FAISS ingestion.
    """

    def __init__(self, chunk_size=500, chunk_overlap=100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    # ---------------------------------------------------------
    # PDF extraction
    # ---------------------------------------------------------
    def _extract_pdf(self, filepath):
        try:
            reader = PdfReader(filepath)
            text = ""

            for page in reader.pages:
                extracted = page.extract_text() or ""
                text += extracted + "\n"

            return text

        except Exception as e:
            raise RuntimeError(f"PDF extraction failed for {filepath}: {e}")

    # ---------------------------------------------------------
    # TXT / MD extraction
    # ---------------------------------------------------------
    def _extract_text_file(self, filepath):
        try:
            return Path(filepath).read_text(encoding="utf-8", errors="ignore")
        except Exception as e:
            raise RuntimeError(f"Text file extraction failed: {filepath}: {e}")

    # ---------------------------------------------------------
    # DOCX extraction
    # ---------------------------------------------------------
    def _extract_docx(self, filepath):
        try:
            import docx
            doc = docx.Document(filepath)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            # Avoid hard errors if python-docx isn't installed
            return ""

    # ---------------------------------------------------------
    # Main processing entry point
    # ---------------------------------------------------------
    def process_file(self, filepath):
        ext = Path(filepath).suffix.lower()

        if ext == ".pdf":
            text = self._extract_pdf(filepath)

        elif ext in [".txt", ".md"]:
            text = self._extract_text_file(filepath)

        elif ext == ".docx":
            text = self._extract_docx(filepath)

        else:
            # Unsupported file type — silently skip
            return []

        # Skip empty docs
        if not text.strip():
            return []

        # Split into document chunks for RAG
        return self._chunk_text(text)



    def _chunk_text(self, text):
        """Return a list of Document objects instead of raw strings."""
        chunks = []
        words = text.split()
        i = 0

        while i < len(words):
            chunk_words = words[i : i + self.chunk_size]
            chunk = " ".join(chunk_words).strip()

            if chunk:
                chunks.append(
                    Document(
                        page_content=chunk,
                        metadata={"source": "pdf"}
                    )
                )

            i += self.chunk_size - self.chunk_overlap

        return chunks

